"""Export functionality for ERP AI Assistant - Excel, CSV, PDF, and Word export."""

import csv
import json
import uuid
from io import BytesIO
from typing import Any, Dict, List

import frappe


def _stringify_cell(value: Any) -> str:
    """Convert a Python value into a clean string for export cells."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, list):
        return ", ".join(str(v) for v in value)
    return str(value)


def _payload_to_rows(payload: Any) -> list[dict[str, Any]]:
    payload = _unwrap_export_payload(payload)
    if payload is None:
        return []
    if isinstance(payload, list):
        rows = []
        for row in payload:
            if isinstance(row, dict):
                rows.append(row)
            else:
                rows.append({"value": _stringify_cell(row)})
        return rows
    if isinstance(payload, dict):
        if isinstance(payload.get("data"), list):
            return _payload_to_rows(payload.get("data"))
        if isinstance(payload.get("result"), list):
            return _payload_to_rows(payload.get("result"))
        compact = {k: v for k, v in payload.items() if k not in {"message", "success"}}
        if compact:
            return [compact]
    return [{"value": _stringify_cell(payload)}]


def _attachment_label_for_filename(file_name: str) -> str:
    extension = str(file_name or "").strip().rsplit(".", 1)
    ext = extension[1].lower() if len(extension) == 2 else ""
    return {
        "xlsx": "Excel",
        "xls": "Excel",
        "csv": "CSV",
        "pdf": "PDF",
        "docx": "Word",
        "doc": "Word",
    }.get(ext, "File")


def _extract_direct_file_attachment(payload: Any, *, max_depth: int = 5) -> dict[str, str] | None:
    def _walk(value: Any, depth: int) -> dict[str, str] | None:
        if depth > max_depth:
            return None
        if isinstance(value, dict):
            file_url = str(value.get("file_url") or "").strip()
            file_name = str(value.get("file_name") or "").strip()
            if file_url and file_name:
                extension = file_name.rsplit(".", 1)
                file_type = extension[1].lower() if len(extension) == 2 else ""
                return {
                    "id": uuid.uuid4().hex,
                    "label": _attachment_label_for_filename(file_name),
                    "filename": file_name,
                    "file_type": file_type,
                    "file_url": file_url,
                }
            for key in ("result", "data", "normalized_result"):
                nested = value.get(key)
                if isinstance(nested, (dict, list)):
                    found = _walk(nested, depth + 1)
                    if found:
                        return found
        elif isinstance(value, list):
            for item in value:
                found = _walk(item, depth + 1)
                if found:
                    return found
        return None

    return _walk(payload, 0)


def _unwrap_export_payload(payload: Any) -> Any:
    current = payload
    for _ in range(5):
        if not isinstance(current, dict):
            return current
        if isinstance(current.get("data"), list):
            return current.get("data")
        if isinstance(current.get("result"), list):
            return current.get("result")
        nested = current.get("result")
        if isinstance(nested, dict) and nested is not current:
            current = nested
            continue
        data_value = current.get("data")
        if isinstance(data_value, dict) and data_value is not current:
            current = data_value
            continue
        break
    return current


def _normalize_export_payload(payload: str) -> tuple[str, list[dict[str, Any]]]:
    data = json.loads(payload)
    title = data.get("title", "export")
    rows = data.get("rows", [])
    if not isinstance(rows, list):
        rows = []
    return title, rows


def _slugify_filename(text: str) -> str:
    """Convert any text into a filesystem-safe filename."""
    import re

    text = re.sub(r"[^\w\s-]", "", text).strip()
    text = re.sub(r"[-\s]+", "-", text)
    return text.lower() or "export"


@frappe.whitelist()
def export_to_excel(payload: str, filename: str | None = None):
    """Export payload data to Excel format."""
    try:
        import openpyxl  # noqa: F401
    except ImportError as exc:
        frappe.throw("Excel export requires openpyxl. Install with: pip install openpyxl")

    title, rows = _normalize_export_payload(payload)

    if not rows:
        frappe.throw("No data to export")

    bytes_content = _build_excel_bytes(title, rows)

    fname = _slugify_filename(filename or title)
    frappe.local.response.filename = f"{fname}.xlsx"
    frappe.local.response.filecontent = bytes_content
    frappe.local.response.type = "download"


@frappe.whitelist()
def export_to_csv(payload: str, filename: str | None = None):
    """Export payload data to CSV format."""
    title, rows = _normalize_export_payload(payload)

    if not rows:
        frappe.throw("No data to export")

    bytes_content = _build_csv_bytes(title, rows)

    fname = _slugify_filename(filename or title)
    frappe.local.response.filename = f"{fname}.csv"
    frappe.local.response.filecontent = bytes_content
    frappe.local.response.type = "download"


@frappe.whitelist()
def export_to_pdf(payload: str, filename: str | None = None):
    """Export payload data to PDF format."""
    try:
        import reportlab  # noqa: F401
    except ImportError as exc:
        frappe.throw("PDF export requires reportlab. Install with: pip install reportlab")

    title, rows = _normalize_export_payload(payload)

    if not rows:
        frappe.throw("No data to export")

    bytes_content = _build_pdf_bytes(title, rows)

    fname = _slugify_filename(filename or title)
    frappe.local.response.filename = f"{fname}.pdf"
    frappe.local.response.filecontent = bytes_content
    frappe.local.response.type = "download"


@frappe.whitelist()
def export_to_word(payload: str, filename: str | None = None):
    """Export payload data to Word format."""
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        frappe.throw("Word export requires python-docx. Install with: pip install python-docx")

    title, rows = _normalize_export_payload(payload)

    if not rows:
        frappe.throw("No data to export")

    bytes_content = _build_word_bytes(title, rows)

    fname = _slugify_filename(filename or title)
    frappe.local.response.filename = f"{fname}.docx"
    frappe.local.response.filecontent = bytes_content
    frappe.local.response.type = "download"


def create_message_artifacts(
    payload: Any,
    title: str,
    formats: list[str] | None = None,
) -> dict[str, Any]:
    direct_attachment = _extract_direct_file_attachment(payload)
    if direct_attachment:
        return {"attachments": [direct_attachment], "exports": {}}

    rows = _payload_to_rows(payload)
    if not rows:
        return {"attachments": [], "exports": {}}

    attachments: list[dict[str, str]] = []
    exports: dict[str, dict[str, Any]] = {}
    base_name = _slugify_filename(title or "assistant-export")
    normalized_title = title or "Assistant Export"
    export_id = uuid.uuid4().hex
    exports[export_id] = {
        "title": normalized_title,
        "rows": rows,
    }

    builders = [
        ("xlsx", "Excel"),
        ("csv", "CSV"),
        ("pdf", "PDF"),
        ("docx", "Word"),
    ]
    allowed_formats = {str(fmt).lower().strip() for fmt in (formats or []) if str(fmt).strip()}
    if allowed_formats:
        builders = [row for row in builders if row[0] in allowed_formats]

    for ext, label in builders:
        attachments.append(
            {
                "id": uuid.uuid4().hex,
                "label": label,
                "filename": f"{base_name}.{ext}",
                "file_type": ext,
                "export_id": export_id,
            }
        )

    return {"attachments": attachments, "exports": exports}


def add_message_attachment_urls(message_name: str, package: dict[str, Any]) -> dict[str, Any]:
    attachments = package.get("attachments") or []
    exports = package.get("exports") or {}
    hydrated = []
    for item in attachments:
        if not isinstance(item, dict):
            continue
        row = dict(item)
        if str(row.get("file_url") or "").strip():
            hydrated.append(row)
            continue
        attachment_id = str(row.get("id") or "").strip()
        export_id = str(row.get("export_id") or "").strip()
        if not attachment_id or export_id not in exports:
            continue
        row["file_url"] = (
            "/api/method/erp_ai_assistant.api.export.download_message_attachment"
            f"?message={message_name}&attachment_id={attachment_id}"
        )
        hydrated.append(row)
    return {"attachments": hydrated, "exports": exports}


def _parse_attachment_package(raw: Any) -> dict[str, Any]:
    if isinstance(raw, str):
        try:
            raw = json.loads(raw)
        except Exception:
            return {"attachments": [], "exports": {}}
    if isinstance(raw, list):
        return {"attachments": raw, "exports": {}}
    if not isinstance(raw, dict):
        return {"attachments": [], "exports": {}}
    attachments = raw.get("attachments")
    exports = raw.get("exports")
    return {
        "attachments": attachments if isinstance(attachments, list) else [],
        "exports": exports if isinstance(exports, dict) else {},
    }


def _export_builder(file_type: str):
    builders = {
        "xlsx": _build_excel_bytes,
        "csv": _build_csv_bytes,
        "pdf": _build_pdf_bytes,
        "docx": _build_word_bytes,
    }
    return builders.get(str(file_type or "").strip().lower())


def _download_permission_check(message_name: str) -> None:
    if not frappe.db.exists("AI Message", message_name):
        raise frappe.DoesNotExistError(f"AI Message {message_name} not found")
    conversation = frappe.db.get_value("AI Message", message_name, "conversation")
    if not conversation:
        raise frappe.PermissionError
    owner = frappe.db.get_value("AI Conversation", conversation, "owner")
    if owner != frappe.session.user and "System Manager" not in frappe.get_roles(frappe.session.user):
        raise frappe.PermissionError


@frappe.whitelist()
def download_message_attachment(message: str, attachment_id: str):
    _download_permission_check(message)
    raw = frappe.db.get_value("AI Message", message, "attachments_json")
    package = _parse_attachment_package(raw)
    attachments = package.get("attachments") or []
    exports = package.get("exports") or {}

    target = None
    for item in attachments:
        if str(item.get("id") or "").strip() == str(attachment_id or "").strip():
            target = item
            break
    if not target:
        frappe.throw("Attachment not found")

    export_id = str(target.get("export_id") or "").strip()
    export_payload = exports.get(export_id)
    if not isinstance(export_payload, dict):
        frappe.throw("Attachment payload not found")

    rows = export_payload.get("rows")
    title = str(export_payload.get("title") or "Assistant Export").strip()
    if not isinstance(rows, list) or not rows:
        frappe.throw("No data to export")

    builder = _export_builder(target.get("file_type"))
    if not builder:
        frappe.throw("Unsupported attachment type")

    bytes_content = builder(title, rows)
    frappe.local.response.filename = str(target.get("filename") or "download")
    frappe.local.response.filecontent = bytes_content
    frappe.local.response.type = "download"


def _build_csv_bytes(title: str, rows: list[dict[str, Any]]) -> bytes:
    headers = list(rows[0].keys())
    buffer = BytesIO()
    text_buffer = []
    import io
    handle = io.StringIO()
    writer = csv.writer(handle)
    writer.writerow(headers)
    for row in rows:
        writer.writerow([_stringify_cell(row.get(header, "")) for header in headers])
    return handle.getvalue().encode("utf-8-sig")


def _build_overview_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not rows:
        return []
    headers = list(rows[0].keys())
    return [
        {"Metric": "Rows", "Value": len(rows)},
        {"Metric": "Columns", "Value": len(headers)},
        {"Metric": "Columns List", "Value": ", ".join(headers)},
    ]


def _fit_worksheet(ws, max_width: int = 48) -> None:
    from openpyxl.styles import Alignment
    for column in ws.columns:
        max_length = 0
        col_cells = list(column)
        for cell in col_cells:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            val = _stringify_cell(cell.value)
            if len(val) > max_length:
                max_length = len(val)
        ws.column_dimensions[col_cells[0].column_letter].width = min(max(max_length + 2, 12), max_width)


def _build_excel_bytes(title: str, rows: list[dict[str, Any]]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "Directory"

    headers = list(rows[0].keys())
    ws.append(headers)
    for row in rows:
        ws.append([_stringify_cell(row.get(header, "")) for header in headers])

    header_fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    alt_fill = PatternFill(fill_type="solid", fgColor="F5F8FC")
    for col_idx, _header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    for row_idx in range(2, ws.max_row + 1):
        if row_idx % 2 == 0:
            for col_idx in range(1, ws.max_column + 1):
                ws.cell(row=row_idx, column=col_idx).fill = alt_fill

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    overview_ws = wb.create_sheet(title="Overview", index=0)
    overview_rows = _build_overview_rows(rows)
    overview_headers = list(overview_rows[0].keys())
    overview_ws.append(overview_headers)
    for item in overview_rows:
        overview_ws.append([_stringify_cell(item.get(key, "")) for key in overview_headers])
    for col_idx in range(1, overview_ws.max_column + 1):
        cell = overview_ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    overview_ws.freeze_panes = "A2"
    _fit_worksheet(overview_ws)
    _fit_worksheet(ws)

    for field in _summary_group_fields(rows)[:3]:
        summary_rows = _build_group_summary(rows, field)
        if not summary_rows:
            continue
        sheet_title = _sheet_title_for_group(field)
        summary_ws = wb.create_sheet(title=sheet_title[:30] or "Summary")
        summary_headers = list(summary_rows[0].keys())
        summary_ws.append(summary_headers)
        for item in summary_rows:
            summary_ws.append([_stringify_cell(item.get(key, "")) for key in summary_headers])
        for col_idx in range(1, summary_ws.max_column + 1):
            cell = summary_ws.cell(row=1, column=col_idx)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
        summary_ws.freeze_panes = "A2"
        summary_ws.auto_filter.ref = summary_ws.dimensions
        _fit_worksheet(summary_ws)

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def _summary_group_fields(rows: list[dict[str, Any]]) -> list[str]:
    if not rows:
        return []
    keys = {key for row in rows if isinstance(row, dict) for key in row.keys()}
    ordered = [
        "territory",
        "payment_terms",
        "department",
        "designation",
        "customer_group",
        "supplier_group",
        "status",
        "company",
        "currency",
    ]
    return [field for field in ordered if field in keys]


def _build_group_summary(rows: list[dict[str, Any]], field: str) -> list[dict[str, Any]]:
    counts: Dict[str, int] = {}
    for row in rows:
        if not isinstance(row, dict):
            continue
        raw_value = row.get(field)
        label = _stringify_cell(raw_value).strip() or "Unspecified"
        counts[label] = counts.get(label, 0) + 1
    total = sum(counts.values())
    if total <= 0:
        return []
    sorted_items = sorted(counts.items(), key=lambda item: (-item[1], item[0].lower()))
    out: list[dict[str, Any]] = []
    for label, count in sorted_items:
        percent = round((count / total) * 100, 2)
        out.append(
            {
                field.replace("_", " ").title(): label,
                "Count": count,
                "Percentage": f"{percent}%",
            }
        )
    return out


def _sheet_title_for_group(field: str) -> str:
    mapping = {
        "territory": "By Territory",
        "payment_terms": "By Payment Terms",
        "department": "By Department",
        "designation": "By Designation",
        "customer_group": "By Customer Group",
        "supplier_group": "By Supplier Group",
        "status": "By Status",
        "company": "By Company",
        "currency": "By Currency",
    }
    return mapping.get(field, f"By {field.replace('_', ' ').title()}")


def _build_pdf_bytes(title: str, rows: list[dict[str, Any]]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    headers = list(rows[0].keys())
    styles = getSampleStyleSheet()
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), title=title, leftMargin=0.4 * inch, rightMargin=0.4 * inch, topMargin=0.45 * inch, bottomMargin=0.45 * inch)
    elements = [Paragraph(title, styles["Title"]), Spacer(1, 8)]

    overview_rows = _build_overview_rows(rows)
    overview_table = Table([[item["Metric"], _stringify_cell(item["Value"])] for item in overview_rows], colWidths=[1.7 * inch, 7.8 * inch])
    overview_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.grey),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.extend([overview_table, Spacer(1, 10)])

    max_columns = min(len(headers), 8)
    visible_headers = headers[:max_columns]
    table_data = [[Paragraph(f"<b>{_stringify_cell(header)}</b>", styles["BodyText"]) for header in visible_headers]]
    for row in rows:
        table_data.append([Paragraph(_stringify_cell(row.get(header, ""))[:500].replace("\n", "<br/>"), styles["BodyText"]) for header in visible_headers])

    page_width = landscape(letter)[0] - doc.leftMargin - doc.rightMargin
    col_width = page_width / max(1, len(visible_headers))
    table = Table(table_data, colWidths=[col_width] * len(visible_headers), repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("FONTSIZE", (0, 1), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("TOPPADDING", (0, 1), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 4),
        ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor("#F8F1E8")]),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(table)
    if len(headers) > max_columns:
        elements.extend([Spacer(1, 10), Paragraph(f"Showing the first {max_columns} of {len(headers)} columns in the PDF layout. Use Excel or CSV for the full wide dataset.", styles["Italic"])])
    doc.build(elements)
    return buffer.getvalue()


def _build_word_bytes(title: str, rows: list[dict[str, Any]]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title, 0)
    for item in _build_overview_rows(rows):
        doc.add_paragraph(f"{item['Metric']}: {_stringify_cell(item['Value'])}")

    headers = list(rows[0].keys())
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid"

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_idx, row in enumerate(rows):
        for col_idx, header in enumerate(headers):
            table.rows[row_idx + 1].cells[col_idx].text = _stringify_cell(row.get(header, ""))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
